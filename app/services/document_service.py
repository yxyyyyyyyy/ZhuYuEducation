from __future__ import annotations

import math
import os
import re
from io import BytesIO
from pathlib import Path

import numpy as np
from sqlalchemy import delete, or_, select

from app.core.database import ClassroomORM, KnowledgeChunkORM, KnowledgeDocumentORM, KnowledgeNodeORM, RetrievalCaseORM, UserORM
from app.domain.models import (
    ImportedDocumentFile,
    KnowledgeDirectoryImportRequest,
    KnowledgeDirectoryImportResponse,
    KnowledgeDocumentImportRequest,
    KnowledgeDocumentView,
    KnowledgeChunkPreview,
    KnowledgeSearchHit,
    KnowledgeSearchRequest,
    RetrievalCaseCreate,
    RetrievalCaseRunResponse,
    RetrievalCaseView,
    RetrievalEvaluationRequest,
    RetrievalEvaluationResponse,
    RetrievalQualityCase,
    RetrievalQualityDashboard,
    RetrievalQualityStrategyMetric,
    RetrievalStrategyEvaluation,
)
from app.repositories.sql_repository import sql_repository
from app.services.dashscope_service import DashScopeEmbeddingService
from app.services.retrieval_service import HybridRetrievalService, RetrievalDocument


class DocumentService:
    def __init__(
        self,
        retrieval_service: HybridRetrievalService,
        dashscope_service: DashScopeEmbeddingService,
    ) -> None:
        self.retrieval_service = retrieval_service
        self.dashscope_service = dashscope_service

    def import_documents(
        self,
        request: KnowledgeDocumentImportRequest,
        user_id: int | None = None,
    ) -> list[KnowledgeDocumentView]:
        imported = []
        with sql_repository.session() as session:
            teacher_scope = self._teacher_scope(session, user_id) if user_id is not None else None
            for item in request.documents:
                school_id, grade_level, subject = self._normalize_document_scope(
                    session,
                    teacher_scope=teacher_scope,
                    school_id=item.school_id,
                    grade_level=item.grade_level,
                    subject=item.subject,
                    topic_id=item.topic_id,
                    require_scope=False,
                )
                doc = KnowledgeDocumentORM(
                    school_id=school_id,
                    teacher_user_id=user_id,
                    grade_level=grade_level,
                    subject=subject,
                    title=item.title,
                    topic_id=item.topic_id,
                    doc_type=item.doc_type,
                    source_name=item.source_name,
                    content=item.content,
                )
                session.add(doc)
                session.flush()
                chunks = self._chunk_text(item.content)
                embeddings = self._embed_chunks(chunks)
                for index, chunk in enumerate(chunks):
                    embedding = embeddings[index] if index < len(embeddings) else None
                    session.add(
                        KnowledgeChunkORM(
                            document_id=doc.id,
                            school_id=school_id,
                            teacher_user_id=user_id,
                            grade_level=grade_level,
                            subject=subject,
                            topic_id=item.topic_id,
                            doc_type=item.doc_type,
                            source_name=item.source_name,
                            chunk_index=index,
                            content=chunk,
                            embedding=embedding,
                            embedding_model=self.dashscope_service.embedding_model if embedding else None,
                            embedding_dim=len(embedding) if embedding else None,
                        )
                    )
                session.flush()
                topic_meta = self._topic_meta_map(session, [doc.topic_id] if doc.topic_id else [])
                imported.append(self._document_view(
                    doc,
                    chunks=[
                        chunk_row
                        for chunk_row in session.execute(
                            select(KnowledgeChunkORM).where(KnowledgeChunkORM.document_id == doc.id)
                        ).scalars().all()
                    ],
                    topic_meta=topic_meta,
                ))
        return imported

    def import_uploaded_document(
        self,
        filename: str,
        content: bytes,
        title: str | None,
        grade_level: str,
        subject: str,
        doc_type: str,
        source_name: str | None,
        user_id: int,
    ) -> KnowledgeDocumentView:
        text = self._read_upload(filename, content)
        if not text.strip():
            raise ValueError("文件已上传但未提取到可读文本")
        request = KnowledgeDocumentImportRequest(
            documents=[
                {
                    "title": (title or Path(filename).stem).strip(),
                    "grade_level": grade_level,
                    "subject": subject,
                    "topic_id": None,
                    "doc_type": doc_type,
                    "source_name": (source_name or filename).strip(),
                    "content": text,
                }
            ]
        )
        return self.import_documents(request, user_id=user_id)[0]

    def list_documents(
        self,
        user_id: int | None = None,
        grade_level: str | None = None,
        subject: str | None = None,
    ) -> list[KnowledgeDocumentView]:
        with sql_repository.session() as session:
            stmt = select(KnowledgeDocumentORM).order_by(KnowledgeDocumentORM.created_at.desc())
            teacher_scope = self._teacher_scope(session, user_id) if user_id is not None else None
            if user_id is not None:
                stmt = stmt.where(KnowledgeDocumentORM.school_id == teacher_scope["school_id"])
            normalized_grade = (grade_level or "").strip()
            normalized_subject = (subject or "").strip()
            if normalized_grade:
                self._validate_scope_choice(session, teacher_scope, normalized_grade, normalized_subject if normalized_subject else None)
                stmt = stmt.where(KnowledgeDocumentORM.grade_level == normalized_grade)
            if normalized_subject:
                stmt = stmt.where(KnowledgeDocumentORM.subject == normalized_subject)
            docs = session.execute(stmt).scalars().all()
            topic_meta = self._topic_meta_map(session, [doc.topic_id for doc in docs if doc.topic_id])
            views = []
            for doc in docs:
                chunks = session.execute(
                    select(KnowledgeChunkORM)
                    .where(KnowledgeChunkORM.document_id == doc.id)
                    .order_by(KnowledgeChunkORM.chunk_index)
                ).scalars().all()
                views.append(self._document_view(doc, chunks, topic_meta))
            return views

    def delete_document(self, document_id: int, user_id: int) -> None:
        with sql_repository.session() as session:
            doc = session.execute(
                select(KnowledgeDocumentORM).where(
                    KnowledgeDocumentORM.id == document_id,
                    KnowledgeDocumentORM.teacher_user_id == user_id,
                )
            ).scalars().first()
            if not doc:
                raise ValueError("document not found")
            session.execute(delete(KnowledgeChunkORM).where(KnowledgeChunkORM.document_id == doc.id))
            session.delete(doc)

    def import_directory(
        self,
        request: KnowledgeDirectoryImportRequest,
        user_id: int | None = None,
    ) -> KnowledgeDirectoryImportResponse:
        directory = Path(request.directory_path).expanduser()
        if not directory.exists() or not directory.is_dir():
            raise ValueError("directory not found")
        allowed_root = os.getenv("DOCUMENT_IMPORT_ROOT", "").strip()
        if allowed_root:
            root = Path(allowed_root).expanduser().resolve()
            resolved_directory = directory.resolve()
            if resolved_directory != root and root not in resolved_directory.parents:
                raise ValueError("directory is outside DOCUMENT_IMPORT_ROOT")

        pattern = "**/*" if request.recursive else "*"
        supported_files = []
        for path in sorted(directory.glob(pattern)):
            if path.is_file() and path.suffix.lower() in {".md", ".markdown", ".txt", ".pdf", ".docx"}:
                supported_files.append(path)
            if len(supported_files) >= request.limit:
                break

        files: list[ImportedDocumentFile] = []
        import_items = []
        for path in supported_files:
            content = self._read_file(path)
            if not content.strip():
                files.append(
                    ImportedDocumentFile(
                        file_path=str(path),
                        title=path.stem,
                        topic_id=request.topic_id or self._guess_topic_id(path.stem),
                        doc_type=request.doc_type or self._guess_doc_type(path.name),
                        imported=False,
                        reason="未解析出有效文本",
                    )
                )
                continue
            topic_id = request.topic_id or self._guess_topic_id(f"{path.stem}\n{content[:120]}")
            doc_type = request.doc_type or self._guess_doc_type(path.name)
            import_items.append(
                {
                    "title": path.stem,
                    "topic_id": topic_id,
                    "doc_type": doc_type,
                    "source_name": str(path.relative_to(directory)),
                    "content": content,
                }
            )
            files.append(
                ImportedDocumentFile(
                    file_path=str(path),
                    title=path.stem,
                    topic_id=topic_id,
                    doc_type=doc_type,
                    imported=True,
                )
            )

        if import_items:
            self.import_documents(KnowledgeDocumentImportRequest(documents=import_items), user_id=user_id)

        imported_count = sum(1 for item in files if item.imported)
        return KnowledgeDirectoryImportResponse(
            imported_count=imported_count,
            skipped_count=len(files) - imported_count,
            files=files,
        )

    def rebuild_embeddings(self, user_id: int | None = None) -> int:
        if not self.dashscope_service.enabled:
            return 0
        with sql_repository.session() as session:
            stmt = select(KnowledgeChunkORM).order_by(KnowledgeChunkORM.id)
            if user_id is not None:
                teacher_scope = self._teacher_scope(session, user_id)
                stmt = stmt.where(KnowledgeChunkORM.school_id == teacher_scope["school_id"])
            chunks = session.execute(stmt).scalars().all()
            texts = [chunk.content for chunk in chunks]
            if not texts:
                return 0
            try:
                embeddings = self.dashscope_service.embed_texts(texts, text_type="document")
            except Exception:
                return 0
            updated = 0
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = embedding
                chunk.embedding_model = self.dashscope_service.embedding_model
                chunk.embedding_dim = len(embedding)
                updated += 1
            return updated

    def search(
        self,
        request: KnowledgeSearchRequest,
        user_id: int | None = None,
    ) -> list[KnowledgeSearchHit]:
        with sql_repository.session() as session:
            teacher_scope = self._teacher_scope(session, user_id) if user_id is not None else None
            normalized_grade = (request.grade_level or "").strip()
            normalized_subject = (request.subject or "").strip()
            self._validate_scope_choice(session, teacher_scope, normalized_grade, normalized_subject, require_subject=True)
            stmt = select(KnowledgeChunkORM)
            if user_id is not None:
                stmt = stmt.where(KnowledgeChunkORM.school_id == teacher_scope["school_id"])
            stmt = stmt.where(KnowledgeChunkORM.grade_level == normalized_grade)
            stmt = stmt.where(KnowledgeChunkORM.subject == normalized_subject)
            chunks = session.execute(stmt).scalars().all()
            doc_stmt = select(KnowledgeDocumentORM)
            if user_id is not None:
                doc_stmt = doc_stmt.where(KnowledgeDocumentORM.school_id == teacher_scope["school_id"])
            docs = {doc.id: doc for doc in session.execute(doc_stmt).scalars().all()}

        preferred_doc_type = self._preferred_doc_type(request.query)
        dense_scores = self._dense_scores(request.query, chunks)
        retrieval_docs = [
            RetrievalDocument(
                identifier=chunk.id,
                title=(docs.get(chunk.document_id).title if docs.get(chunk.document_id) else "未知文档"),
                content=chunk.content,
                topic_id=chunk.topic_id,
                doc_type=chunk.doc_type,
                source_name=chunk.source_name,
            )
            for chunk in chunks
        ]
        chunk_map = {chunk.id: chunk for chunk in chunks}
        ranked = self.retrieval_service.rank_with_strategy(
            strategy=request.strategy,
            query=request.query,
            documents=retrieval_docs,
            topic_id=request.topic_id,
            preferred_doc_type=preferred_doc_type,
            dense_scores=dense_scores,
            limit=request.limit,
        )
        rerank_scores = self._rerank_scores(request.query, ranked)
        if rerank_scores and request.strategy == "rerank":
            ranked = self.retrieval_service.rank_with_strategy(
                strategy=request.strategy,
                query=request.query,
                documents=[
                    RetrievalDocument(
                        identifier=item.identifier,
                        title=item.title,
                        content=item.content,
                        topic_id=item.topic_id,
                        doc_type=item.doc_type,
                        source_name=item.source_name,
                    )
                    for item in ranked
                ],
                topic_id=request.topic_id,
                preferred_doc_type=preferred_doc_type,
                dense_scores={item.identifier: dense_scores.get(item.identifier, 0.0) for item in ranked},
                rerank_scores=rerank_scores,
                limit=request.limit,
            )
        return [
            KnowledgeSearchHit(
                document_title=item.title,
                doc_type=item.doc_type or "unknown",
                source_name=item.source_name or "unknown",
                grade_level=(chunk_map.get(item.identifier).grade_level if chunk_map.get(item.identifier) else ""),
                subject=(chunk_map.get(item.identifier).subject if chunk_map.get(item.identifier) else ""),
                topic_id=item.topic_id,
                snippet=item.content[:160],
                score=item.final_score,
                lexical_score=item.lexical_score,
                vector_score=item.vector_score,
                dense_score=item.dense_score,
                rerank_score=item.rerank_score,
            )
            for item in ranked
        ]

    def evaluate_retrieval(
        self,
        request: RetrievalEvaluationRequest,
        user_id: int | None = None,
    ) -> RetrievalEvaluationResponse:
        strategies = []
        for strategy in ["keyword", "hybrid", "dense", "rerank"]:
            hits = self.search(
                KnowledgeSearchRequest(
                    query=request.query,
                    topic_id=request.topic_id,
                    limit=request.limit,
                    strategy=strategy,
                ),
                user_id=user_id,
            )
            hit_at_1 = bool(hits[:1] and self._matches_expected(hits[0], request.expected_topic_id, request.expected_doc_type))
            hit_at_3 = any(
                self._matches_expected(item, request.expected_topic_id, request.expected_doc_type)
                for item in hits[:3]
            )
            mrr = 0.0
            for index, item in enumerate(hits, start=1):
                if self._matches_expected(item, request.expected_topic_id, request.expected_doc_type):
                    mrr = round(1 / index, 4)
                    break
            strategies.append(
                RetrievalStrategyEvaluation(
                    strategy=strategy,
                    hit_at_1=hit_at_1,
                    hit_at_3=hit_at_3,
                    mrr=mrr,
                    hits=hits,
                )
            )
        best_strategy = max(strategies, key=lambda item: (item.hit_at_1, item.hit_at_3, item.mrr)).strategy
        return RetrievalEvaluationResponse(
            query=request.query,
            expected_topic_id=request.expected_topic_id,
            expected_doc_type=request.expected_doc_type,
            best_strategy=best_strategy,
            strategies=strategies,
        )

    def retrieval_quality_dashboard(self, user_id: int | None = None) -> RetrievalQualityDashboard:
        cases = [
            ("教材理解函数对应关系", "教材里怎么理解函数对应关系", "functions", "textbook"),
            ("讲义理解斜率截距", "讲义里一次函数的斜率和截距", "linear_functions", "handout"),
            ("题解分析函数概念题", "题解中怎么分析函数概念题", "functions", "solution"),
            ("讲义判断 y=kx+b", "讲义里如何根据 y=kx+b 判断截距", "linear_functions", "handout"),
        ]
        case_results = []
        metric_rows = {name: {"hit1": 0, "hit3": 0, "mrr": 0.0} for name in ["keyword", "hybrid", "dense", "rerank"]}

        for label, query, topic_id, doc_type in cases:
            result = self.evaluate_retrieval(
                RetrievalEvaluationRequest(
                    query=query,
                    topic_id=topic_id,
                    expected_topic_id=topic_id,
                    expected_doc_type=doc_type,
                    limit=5,
                ),
                user_id=user_id,
            )
            case_results.append(
                RetrievalQualityCase(
                    label=label,
                    query=query,
                    expected_topic_id=topic_id,
                    expected_doc_type=doc_type,
                    best_strategy=result.best_strategy,
                    strategies=result.strategies,
                )
            )
            for item in result.strategies:
                metric_rows[item.strategy]["hit1"] += 1 if item.hit_at_1 else 0
                metric_rows[item.strategy]["hit3"] += 1 if item.hit_at_3 else 0
                metric_rows[item.strategy]["mrr"] += item.mrr

        total = len(cases) or 1
        metrics = [
            RetrievalQualityStrategyMetric(
                strategy=strategy,
                hit_at_1=round(values["hit1"] / total, 2),
                hit_at_3=round(values["hit3"] / total, 2),
                mrr=round(values["mrr"] / total, 2),
            )
            for strategy, values in metric_rows.items()
        ]
        metrics.sort(key=lambda item: (item.hit_at_1, item.hit_at_3, item.mrr), reverse=True)
        return RetrievalQualityDashboard(
            total_cases=total,
            strategies=metrics,
            cases=case_results,
        )

    def list_retrieval_cases(self, user_id: int) -> list[RetrievalCaseView]:
        with sql_repository.session() as session:
            rows = session.execute(
                select(RetrievalCaseORM)
                .where(RetrievalCaseORM.teacher_user_id == user_id)
                .order_by(RetrievalCaseORM.created_at.desc())
            ).scalars().all()
            return [self._retrieval_case_view(row) for row in rows]

    def create_retrieval_case(self, user_id: int, request: RetrievalCaseCreate) -> RetrievalCaseView:
        with sql_repository.session() as session:
            row = RetrievalCaseORM(
                teacher_user_id=user_id,
                label=request.label,
                query=request.query,
                expected_topic_id=request.expected_topic_id,
                expected_doc_type=request.expected_doc_type,
            )
            session.add(row)
            session.flush()
            return self._retrieval_case_view(row)

    def delete_retrieval_case(self, user_id: int, case_id: int) -> None:
        with sql_repository.session() as session:
            row = session.execute(
                select(RetrievalCaseORM).where(
                    RetrievalCaseORM.id == case_id,
                    RetrievalCaseORM.teacher_user_id == user_id,
                )
            ).scalars().first()
            if not row:
                raise ValueError("retrieval case not found")
            session.delete(row)

    def run_retrieval_cases(self, user_id: int) -> RetrievalCaseRunResponse:
        cases = self.list_retrieval_cases(user_id)
        case_results = []
        hit1_total = 0
        hit3_total = 0
        mrr_total = 0.0
        for case in cases:
            result = self.evaluate_retrieval(
                RetrievalEvaluationRequest(
                    query=case.query,
                    topic_id=case.expected_topic_id,
                    expected_topic_id=case.expected_topic_id,
                    expected_doc_type=case.expected_doc_type,
                    limit=5,
                ),
                user_id=user_id,
            )
            best = max(result.strategies, key=lambda item: (item.hit_at_1, item.hit_at_3, item.mrr))
            hit1_total += 1 if best.hit_at_1 else 0
            hit3_total += 1 if best.hit_at_3 else 0
            mrr_total += best.mrr
            case_results.append(
                RetrievalQualityCase(
                    label=case.label,
                    query=case.query,
                    expected_topic_id=case.expected_topic_id,
                    expected_doc_type=case.expected_doc_type,
                    best_strategy=result.best_strategy,
                    strategies=result.strategies,
                )
            )
        total = len(cases)
        divisor = total or 1
        return RetrievalCaseRunResponse(
            total_cases=total,
            hit_at_1=round(hit1_total / divisor, 2),
            hit_at_3=round(hit3_total / divisor, 2),
            mrr=round(mrr_total / divisor, 2),
            cases=case_results,
        )

    def _chunk_text(self, text: str, chunk_size: int = 160) -> list[str]:
        compact = text.strip()
        if not compact:
            return []
        return [compact[index : index + chunk_size] for index in range(0, len(compact), chunk_size)]

    def _document_view(
        self,
        doc: KnowledgeDocumentORM,
        chunks: list[KnowledgeChunkORM],
        topic_meta: dict[str, dict] | None = None,
    ) -> KnowledgeDocumentView:
        topic_meta = topic_meta or {}
        meta = topic_meta.get(doc.topic_id or "", {})
        previews = [
            KnowledgeChunkPreview(
                id=chunk.id,
                chunk_index=chunk.chunk_index,
                content=chunk.content[:220],
                embedding_ready=bool(chunk.embedding),
            )
            for chunk in chunks[:3]
        ]
        return KnowledgeDocumentView(
            id=doc.id,
            title=doc.title,
            school_id=doc.school_id,
            topic_id=doc.topic_id,
            subject=doc.subject or meta.get("subject", ""),
            grade_level=doc.grade_level or meta.get("grade_level", ""),
            doc_type=doc.doc_type,
            source_name=doc.source_name,
            chunk_count=len(chunks),
            embedding_ready_count=sum(1 for chunk in chunks if chunk.embedding),
            created_at=doc.created_at,
            content_preview=(doc.content or "")[:220],
            chunk_previews=previews,
            can_delete=bool(doc.teacher_user_id),
        )

    def _teacher_scope(self, session, user_id: int | None) -> dict | None:
        if user_id is None:
            return None
        teacher = session.execute(select(UserORM).where(UserORM.id == user_id)).scalars().first()
        if not teacher or not teacher.school_id:
            raise ValueError("当前老师未绑定学校")
        grades = sorted({
            grade.strip()
            for grade in session.execute(
                select(ClassroomORM.grade_level).where(ClassroomORM.teacher_user_id == user_id)
            ).scalars().all()
            if grade and grade.strip()
        })
        return {
            "school_id": teacher.school_id,
            "teacher_user_id": user_id,
            "grades": grades,
        }

    def _normalize_document_scope(
        self,
        session,
        *,
        teacher_scope: dict | None,
        school_id: int | None,
        grade_level: str | None,
        subject: str | None,
        topic_id: str | None,
        require_scope: bool,
    ) -> tuple[int | None, str, str]:
        normalized_school_id = school_id
        normalized_grade = (grade_level or "").strip()
        normalized_subject = (subject or "").strip()
        if topic_id:
            topic_meta = self._topic_meta_map(session, [topic_id]).get(topic_id, {})
            if not normalized_grade:
                normalized_grade = (topic_meta.get("grade_level") or "").strip()
            if not normalized_subject:
                normalized_subject = (topic_meta.get("subject") or "").strip()
        if teacher_scope is not None:
            normalized_school_id = teacher_scope["school_id"]
            self._validate_scope_choice(
                session,
                teacher_scope,
                normalized_grade,
                normalized_subject if normalized_subject else None,
                require_subject=require_scope,
            )
        elif require_scope and (not normalized_grade or not normalized_subject):
            raise ValueError("资料需要绑定年级和学科")
        return normalized_school_id, normalized_grade, normalized_subject

    def _validate_scope_choice(
        self,
        session,
        teacher_scope: dict | None,
        grade_level: str | None,
        subject: str | None = None,
        require_subject: bool = False,
    ) -> None:
        normalized_grade = (grade_level or "").strip()
        normalized_subject = (subject or "").strip()
        if not normalized_grade:
            raise ValueError("请选择年级")
        if require_subject and not normalized_subject:
            raise ValueError("请选择学科")
        if teacher_scope is None:
            return
        allowed_grades = teacher_scope.get("grades", [])
        if allowed_grades and normalized_grade not in allowed_grades:
            raise ValueError("当前年级不在老师所教班级范围内")
        if normalized_subject:
            row = session.execute(
                select(KnowledgeNodeORM.id).where(
                    KnowledgeNodeORM.is_deleted == 0,
                    KnowledgeNodeORM.school_id == teacher_scope["school_id"],
                    KnowledgeNodeORM.grade_level == normalized_grade,
                    KnowledgeNodeORM.subject == normalized_subject,
                )
            ).scalars().first()
            if not row:
                raise ValueError("当前学科不在老师可用知识范围内")

    def _topic_meta_map(self, session, topic_ids: list[str]) -> dict[str, dict]:
        normalized_ids = {topic_id.strip() for topic_id in topic_ids if topic_id and topic_id.strip()}
        if not normalized_ids:
            return {}
        rows = session.execute(
            select(KnowledgeNodeORM).where(
                KnowledgeNodeORM.is_deleted == 0,
                or_(KnowledgeNodeORM.node_key.in_(normalized_ids), KnowledgeNodeORM.topic_ref_id.in_(normalized_ids)),
            )
        ).scalars().all()
        meta_map: dict[str, dict] = {}
        for row in rows:
            meta = {
                "subject": row.subject or "",
                "grade_level": row.grade_level or "",
                "name": row.name,
            }
            meta_map[row.node_key] = meta
            if row.topic_ref_id:
                meta_map[row.topic_ref_id] = meta
        return meta_map

    def _retrieval_case_view(self, row: RetrievalCaseORM) -> RetrievalCaseView:
        return RetrievalCaseView(
            id=row.id,
            label=row.label,
            query=row.query,
            expected_topic_id=row.expected_topic_id,
            expected_doc_type=row.expected_doc_type,
            created_at=row.created_at,
        )

    def _preferred_doc_type(self, query: str) -> str | None:
        lowered = query.lower()
        if "教材" in lowered:
            return "textbook"
        if "讲义" in lowered:
            return "handout"
        if "题解" in lowered or "解析" in lowered:
            return "solution"
        return None

    def _guess_doc_type(self, filename: str) -> str:
        lowered = filename.lower()
        if "教材" in lowered or "textbook" in lowered:
            return "textbook"
        if "讲义" in lowered or "handout" in lowered or "notes" in lowered:
            return "handout"
        if "题解" in lowered or "解析" in lowered or "solution" in lowered:
            return "solution"
        return "reference"

    def _guess_topic_id(self, text: str) -> str | None:
        mapping = {
            "linear_functions": ["一次函数", "斜率", "截距", "y=kx+b", "linear"],
            "functions": ["函数", "自变量", "因变量", "对应关系", "function"],
            "equations": ["方程", "移项", "解方程", "equation"],
            "arithmetic": ["算术", "四则", "运算", "arithmetic"],
        }
        lowered = text.lower()
        for topic_id, keywords in mapping.items():
            if any(keyword.lower() in lowered for keyword in keywords):
                return topic_id
        return None

    def _read_file(self, path: Path) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix in {".md", ".markdown", ".txt"}:
                return path.read_text(encoding="utf-8", errors="ignore")
            if suffix == ".pdf":
                content = path.read_bytes()
                extracted = self._read_pdf_text_bytes(content)
                if self._has_readable_text(extracted):
                    return extracted
                ocr_text = self._ocr_pdf_bytes(content)
                if self._has_readable_text(ocr_text):
                    return ocr_text
                return ""
            if suffix == ".docx":
                from docx import Document

                document = Document(str(path))
                return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return ""
        return ""

    def _read_upload(self, filename: str, content: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        try:
            if suffix in {".md", ".markdown", ".txt"}:
                return content.decode("utf-8", errors="ignore")
            if suffix == ".pdf":
                extracted = self._read_pdf_text_bytes(content)
                if self._has_readable_text(extracted):
                    return extracted
                ocr_text = self._ocr_pdf_bytes(content)
                if self._has_readable_text(ocr_text):
                    return ocr_text
                raise ValueError("PDF 未提取到可读文本，OCR 识别失败，请上传文字版 PDF 或 docx")
            if suffix == ".docx":
                from docx import Document

                document = Document(BytesIO(content))
                return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except ValueError:
            raise
        except Exception as exc:
            if suffix == ".pdf":
                raise ValueError("PDF 文件解析失败，请确认文件未损坏后重试") from exc
            if suffix == ".docx":
                raise ValueError("DOCX 文件解析失败，请确认文件格式正确") from exc
            if suffix in {".md", ".markdown", ".txt"}:
                raise ValueError("文本文件读取失败，请确认编码格式正确") from exc
            raise
        raise ValueError("不支持的文件类型，请上传 txt、md、pdf 或 docx")

    def _read_pdf_text_bytes(self, content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    def _has_readable_text(self, text: str, min_chars: int = 12) -> bool:
        readable = "".join(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", text or ""))
        return len(readable) >= min_chars

    def _ocr_pdf_bytes(self, content: bytes) -> str:
        try:
            import fitz
            from PIL import Image
            import pytesseract
        except Exception as exc:
            raise ValueError("PDF 未提取到可读文本，当前服务未启用 OCR，请上传文字版 PDF 或 docx") from exc

        tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        elif Path("/opt/homebrew/bin/tesseract").exists():
            pytesseract.pytesseract.tesseract_cmd = "/opt/homebrew/bin/tesseract"

        try:
            document = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise ValueError("PDF 文件无法解析，请确认文件未损坏后重试") from exc

        ocr_pages: list[str] = []
        try:
            for page in document:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                page_text = ""
                for lang in ("chi_sim+eng", "eng"):
                    try:
                        page_text = pytesseract.image_to_string(image, lang=lang).strip()
                    except Exception:
                        page_text = ""
                    if self._has_readable_text(page_text, min_chars=6):
                        break
                if page_text:
                    ocr_pages.append(page_text)
        finally:
            document.close()

        return "\n".join(ocr_pages).strip()

    def _matches_expected(
        self,
        hit: KnowledgeSearchHit,
        expected_topic_id: str | None,
        expected_doc_type: str | None,
    ) -> bool:
        topic_ok = expected_topic_id is None or hit.topic_id == expected_topic_id
        type_ok = expected_doc_type is None or hit.doc_type == expected_doc_type
        return topic_ok and type_ok

    def _embed_chunks(self, chunks: list[str]) -> list[list[float]]:
        if not chunks:
            return []
        if not self.dashscope_service.enabled:
            return []
        try:
            return self.dashscope_service.embed_texts(chunks, text_type="document")
        except Exception:
            return []

    def _dense_scores(self, query: str, chunks: list[KnowledgeChunkORM]) -> dict[int, float]:
        if not self.dashscope_service.enabled:
            return {}
        vector_chunks = [chunk for chunk in chunks if chunk.embedding]
        if not vector_chunks:
            return {}
        try:
            query_embedding = self.dashscope_service.embed_texts([query], text_type="query")[0]
        except Exception:
            return {}
        query_vector = np.array(query_embedding, dtype=float)
        query_norm = np.linalg.norm(query_vector) or 1.0
        scores = {}
        for chunk in vector_chunks:
            doc_vector = np.array(chunk.embedding, dtype=float)
            doc_norm = np.linalg.norm(doc_vector) or 1.0
            score = float(np.dot(query_vector, doc_vector) / (query_norm * doc_norm))
            scores[chunk.id] = round(max(score, 0.0), 4)
        return scores

    def _rerank_scores(self, query: str, ranked_hits) -> dict[int, float]:
        if not self.dashscope_service.enabled or not ranked_hits:
            return {}
        contents = [item.content for item in ranked_hits[:12]]
        identifiers = [item.identifier for item in ranked_hits[:12]]
        try:
            raw_scores = self.dashscope_service.rerank(query=query, documents=contents, top_n=len(contents))
        except Exception:
            return {}
        if not raw_scores:
            return {}
        max_score = max(raw_scores) or 1.0
        return {
            identifier: round(score / max_score, 4)
            for identifier, score in zip(identifiers, raw_scores)
        }
